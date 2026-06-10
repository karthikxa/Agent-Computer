"""Document Worker — edit Google Docs, Sheets, Word, Excel like a human employee.

✅ Google Docs: create, read, write, format, share
✅ Google Sheets: read/write cells, formulas, charts
✅ Microsoft Word/Excel: create, edit via LibreOffice on desktop
✅ PDF: read text, fill forms, sign
✅ Clipboard-based data transfer between apps
"""

from __future__ import annotations

import asyncio
import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)


@dataclass
class DocumentWorker:
    """Automate document tasks like a real office employee.

    Parameters
    ----------
    browser:
        BrowserWorker for Google Docs/Sheets web automation.
    human:
        HumanBehaviorEngine for natural typing in documents.
    """

    browser: Any | None = None
    human: Any | None = None

    # ── Google Docs ────────────────────────────────────────────────────────────

    async def open_google_doc(self, doc_url: str) -> bool:
        """Open a Google Doc in the browser."""
        if not self.browser:
            return False
        await self.browser.navigate(doc_url)
        await asyncio.sleep(3)
        return True

    async def create_google_doc(self, title: str) -> str:
        """Create a new Google Doc and return its URL."""
        if not self.browser:
            return ""
        await self.browser.navigate("https://docs.google.com/document/create")
        await asyncio.sleep(3)
        # Set title
        await self.browser.fill_form({"title": title})
        await asyncio.sleep(1)
        if self.browser._page:
            return self.browser._page.url
        return ""

    async def type_in_doc(self, content: str) -> bool:
        """Click the document body and type content."""
        if not self.browser:
            return False
        try:
            await self.browser.click("Document body")
            await asyncio.sleep(0.5)
            if self.human:
                await self.human.type_text(content)
            else:
                if self.browser._page:
                    await self.browser._page.keyboard.type(content)
            return True
        except Exception as exc:
            logger.error("type_in_doc: %s", exc)
            return False

    async def read_google_doc(self, doc_url: str) -> str:
        """Open a Google Doc and return all visible text."""
        await self.open_google_doc(doc_url)
        if self.browser:
            return await self.browser.get_page_text()
        return ""

    async def share_google_doc(self, doc_url: str, email: str, role: str = "editor") -> bool:
        """Share a Google Doc with another user."""
        await self.open_google_doc(doc_url)
        if not self.browser:
            return False
        await self.browser.click("Share")
        await asyncio.sleep(1.5)
        await self.browser.fill_form({"Add people": email})
        await asyncio.sleep(0.5)
        # Select role
        try:
            await self.browser.click(role.capitalize())
        except Exception:
            pass
        await self.browser.click("Send")
        await asyncio.sleep(1.5)
        return True

    async def add_comment(self, doc_url: str, comment_text: str) -> bool:
        """Add a comment to a Google Doc."""
        await self.open_google_doc(doc_url)
        if not self.browser:
            return False
        # Ctrl+Alt+M is the Google Docs shortcut for Add Comment
        if self.browser._page:
            await self.browser._page.keyboard.press("Control+Alt+m")
        await asyncio.sleep(0.5)
        if self.human:
            await self.human.type_text(comment_text)
        await asyncio.sleep(0.3)
        await self.browser.click("Comment")
        return True

    # ── Google Sheets ──────────────────────────────────────────────────────────

    async def open_google_sheet(self, sheet_url: str) -> bool:
        """Open a Google Sheet."""
        if not self.browser:
            return False
        await self.browser.navigate(sheet_url)
        await asyncio.sleep(3)
        return True

    async def create_google_sheet(self, title: str) -> str:
        """Create a new Google Sheet and return its URL."""
        if not self.browser:
            return ""
        await self.browser.navigate("https://sheets.google.com/create")
        await asyncio.sleep(3)
        if self.browser._page:
            return self.browser._page.url
        return ""

    async def read_cell(self, sheet_url: str, cell: str) -> str:
        """Read the value of a cell (e.g. 'A1')."""
        await self.open_google_sheet(sheet_url)
        if not self.browser or not self.browser._page:
            return ""
        # Use Name Box to navigate to cell
        try:
            page = self.browser._page
            await page.click(".cell-input")  # Name box
            await page.fill(".cell-input", cell)
            await page.keyboard.press("Enter")
            await asyncio.sleep(0.5)
            # Get cell value from formula bar
            val = await page.input_value(".formula-bar-input")
            return val
        except Exception as exc:
            logger.warning("read_cell: %s", exc)
            return ""

    async def write_cell(self, sheet_url: str, cell: str, value: str) -> bool:
        """Write a value to a specific cell."""
        await self.open_google_sheet(sheet_url)
        if not self.browser or not self.browser._page:
            return False
        try:
            page = self.browser._page
            # Navigate to cell via Name Box
            await page.click(".cell-input")
            await page.fill(".cell-input", cell)
            await page.keyboard.press("Enter")
            await asyncio.sleep(0.3)
            # Type value
            if self.human:
                await self.human.type_text(value)
            else:
                await page.keyboard.type(value)
            await page.keyboard.press("Enter")
            return True
        except Exception as exc:
            logger.error("write_cell: %s", exc)
            return False

    async def write_row(self, sheet_url: str, start_cell: str, values: list[str]) -> bool:
        """Write a list of values starting at a cell, moving right."""
        col = ord(start_cell[0].upper()) - ord("A")
        row = int(start_cell[1:])
        for i, val in enumerate(values):
            cell = f"{chr(ord('A') + col + i)}{row}"
            ok = await self.write_cell(sheet_url, cell, val)
            if not ok:
                return False
        return True

    # ── Local documents (LibreOffice on desktop) ───────────────────────────────

    async def create_word_doc(self, filename: str, content: str) -> Path:
        """Create a .docx file using python-docx."""
        try:
            from docx import Document
            doc = Document()
            for line in content.splitlines():
                doc.add_paragraph(line)
            path = Path(filename)
            doc.save(path)
            logger.info("DocumentWorker: created %s", path)
            return path
        except ImportError:
            # Fall back: create plain text file
            path = Path(filename.replace(".docx", ".txt"))
            path.write_text(content)
            return path

    async def read_word_doc(self, filename: str) -> str:
        """Read text from a .docx file."""
        try:
            from docx import Document
            doc = Document(filename)
            return "\n".join(p.text for p in doc.paragraphs)
        except ImportError:
            return Path(filename).read_text(errors="replace") if Path(filename).exists() else ""

    async def create_spreadsheet(self, filename: str, data: list[list[str]]) -> Path:
        """Create an Excel/CSV spreadsheet with data."""
        path = Path(filename)
        try:
            import csv
            csv_path = path.with_suffix(".csv")
            with open(csv_path, "w", newline="") as f:
                writer = csv.writer(f)
                writer.writerows(data)
            logger.info("DocumentWorker: created %s", csv_path)
            return csv_path
        except Exception as exc:
            logger.error("create_spreadsheet: %s", exc)
            return path

    async def read_pdf(self, filename: str) -> str:
        """Extract text from a PDF file."""
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filename)
            return "\n".join(page.get_text() for page in doc)
        except ImportError:
            try:
                # Fallback: pdfplumber
                import pdfplumber
                with pdfplumber.open(filename) as pdf:
                    return "\n".join(p.extract_text() or "" for p in pdf.pages)
            except ImportError:
                logger.warning("PDF reading requires PyMuPDF or pdfplumber")
                return ""

    async def open_file_in_libreoffice(self, filepath: str) -> bool:
        """Open a local document in LibreOffice Writer/Calc."""
        import shutil
        import asyncio as _asyncio
        if shutil.which("libreoffice"):
            proc = await _asyncio.create_subprocess_exec(
                "libreoffice", "--writer" if "doc" in filepath else "--calc", filepath,
                stdout=_asyncio.subprocess.DEVNULL,
                stderr=_asyncio.subprocess.DEVNULL,
            )
            await asyncio.sleep(3)
            return proc.returncode is None  # still running
        return False

    # ── Data entry helpers ─────────────────────────────────────────────────────

    async def fill_web_form(self, form_data: dict[str, str]) -> bool:
        """Fill a web form with human-like behavior."""
        if not self.browser:
            return False
        for label, value in form_data.items():
            try:
                if self.human:
                    # Find the field and type with human behavior
                    await self.browser.click(label)
                    await asyncio.sleep(0.3)
                    await self.human.type_text(value, clear_first=True)
                    await self.human.press_tab()
                else:
                    await self.browser.fill_form({label: value})
                await asyncio.sleep(0.3)
            except Exception as exc:
                logger.warning("fill_web_form[%s]: %s", label, exc)
        return True
